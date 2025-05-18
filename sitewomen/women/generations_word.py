from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
from .models import Contest, Application, Review
from django.db.models import Avg
import io

def generate_contest_report(contest_id):
    contest = Contest.objects.get(id=contest_id)
    applications = Application.objects.filter(contest=contest)
    
    # Создаем документ
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Заголовок
    title = doc.add_paragraph(f'Отчет по конкурсу "{contest.title}"')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    
    # Шапка документа
    doc.add_paragraph(f'Название конкурса: {contest.title}')
    doc.add_paragraph(f'Дата проведения конкурса: {contest.start_date} - {contest.end_date}')
    doc.add_paragraph(f'Период отчетности: {contest.start_date} - {contest.end_date}')
    
    # Описание конкурса
    doc.add_paragraph('Описание конкурса:', style='Heading 2')
    doc.add_paragraph(contest.description)
    
    # Результаты конкурса
    doc.add_paragraph('Результаты конкурса:', style='Heading 2')
    
    # Список победителей
    winners = applications.annotate(
        avg_score=Avg('reviews__total_score')
    ).order_by('-avg_score')[:10]
    
    table_winners = doc.add_table(rows=1, cols=4)
    table_winners.style = 'Table Grid'
    hdr_cells = table_winners.rows[0].cells
    hdr_cells[0].text = 'Место'
    hdr_cells[1].text = 'Участник'
    hdr_cells[2].text = 'Работа'
    hdr_cells[3].text = 'Баллы'
    
    for i, app in enumerate(winners, 1):
        row_cells = table_winners.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = f"{app.user.first_name} {app.user.last_name}"
        row_cells[2].text = app.answers.filter(field__name='work_title').first().value if app.answers.filter(field__name='work_title').exists() else 'Нет названия'
        row_cells[3].text = str(app.avg_score or 0)
    
    # Анализ работ участников
    doc.add_paragraph('Анализ работ участников:', style='Heading 2')
    table_analysis = doc.add_table(rows=1, cols=4)
    table_analysis.style = 'Table Grid'
    hdr_cells = table_analysis.rows[0].cells
    hdr_cells[0].text = 'Участник'
    hdr_cells[1].text = 'Работа'
    hdr_cells[2].text = 'Баллы'
    hdr_cells[3].text = 'Комментарии'
    
    for app in applications:
        review = app.reviews.first()
        row_cells = table_analysis.add_row().cells
        row_cells[0].text = f"{app.user.first_name} {app.user.last_name}"
        row_cells[1].text = app.answers.filter(field__name='work_title').first().value if app.answers.filter(field__name='work_title').exists() else 'Нет названия'
        row_cells[2].text = str(review.total_score if review else 0)
        row_cells[3].text = review.comment if review else 'Нет комментариев'
    
    # Статистика
    doc.add_paragraph('Статистика:', style='Heading 2')
    doc.add_paragraph(f'Количество поданных заявок: {applications.count()}')
    
    # Приложения
    doc.add_paragraph('Приложения:', style='Heading 2')
    
    # Список участников
    doc.add_paragraph('Список участников:', style='Heading 3')
    table_participants = doc.add_table(rows=1, cols=3)
    table_participants.style = 'Table Grid'
    hdr_cells = table_participants.rows[0].cells
    hdr_cells[0].text = 'Участник'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Контактная информация'
    
    for app in applications:
        row_cells = table_participants.add_row().cells
        row_cells[0].text = f"{app.user.first_name} {app.user.last_name}"
        row_cells[1].text = app.user.email
        contact_info = app.answers.filter(field__name='contact_info').first()
        row_cells[2].text = contact_info.value if contact_info else 'Нет контактной информации'
    
    # Сохраняем документ в поток
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    # Создаем HTTP-ответ
    response = HttpResponse(
        doc_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=report_{contest.title}.docx'
    
    return response 